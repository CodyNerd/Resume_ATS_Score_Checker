"""
Text extraction utilities for resume files
Supports PDF, DOCX, and TXT formats
"""

import io
import re
from typing import Union
import streamlit as st

def extract_text(uploaded_file) -> str:
    """
    Extract text from uploaded resume file
    
    Args:
        uploaded_file: Streamlit uploaded file object
    
    Returns:
        str: Extracted text content
    
    Raises:
        ValueError: If file format is not supported or libraries are missing
        Exception: For other processing errors
    """
    
    if uploaded_file is None:
        raise ValueError("No file uploaded")
    
    file_extension = uploaded_file.name.lower().split('.')[-1]
    
    try:
        if file_extension == 'txt':
            return extract_text_from_txt(uploaded_file)
        elif file_extension == 'pdf':
            return extract_text_from_pdf(uploaded_file)
        elif file_extension in ['docx', 'doc']:
            return extract_text_from_docx(uploaded_file)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    except Exception as e:
        raise Exception(f"Error processing {file_extension.upper()} file: {str(e)}")

def extract_text_from_txt(uploaded_file) -> str:
    """
    Extract text from TXT file
    
    Args:
        uploaded_file: Streamlit uploaded file object
    
    Returns:
        str: Text content
    """
    try:
        # Try UTF-8 first
        content = uploaded_file.read().decode('utf-8')
    except UnicodeDecodeError:
        # Fallback to latin-1 if UTF-8 fails
        uploaded_file.seek(0)
        content = uploaded_file.read().decode('latin-1', errors='ignore')
    
    return clean_text(content)

def extract_text_from_pdf(uploaded_file) -> str:
    """
    Extract text from PDF file using pypdf
    
    Args:
        uploaded_file: Streamlit uploaded file object
    
    Returns:
        str: Extracted text content
    
    Raises:
        ValueError: If pypdf is not available
    """
    try:
        from pypdf import PdfReader
        print(f"DEBUG: pypdf imported successfully")
    except ImportError:
        try:
            import PyPDF2
            from PyPDF2 import PdfReader
            print(f"DEBUG: PyPDF2 imported as fallback, version: {PyPDF2.__version__}")
        except ImportError as e:
            print(f"DEBUG: Both pypdf and PyPDF2 import failed: {e}")
            raise ValueError("pypdf or PyPDF2 library is required for PDF processing. Please install with: pip install pypdf")
    
    try:
        # Create a PDF reader object
        pdf_reader = PdfReader(io.BytesIO(uploaded_file.read()))
        
        text_content = []
        
        # Extract text from each page
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text_content.append(page_text)
            print(f"DEBUG: Page {page_num + 1}: {len(page_text)} characters")
        
        # Combine all pages
        full_text = '\n'.join(text_content)
        print(f"DEBUG: Total extracted text: {len(full_text)} characters")
        
        if not full_text.strip():
            raise Exception("No text could be extracted from the PDF. The file might be image-based or corrupted.")
        
        return clean_text(full_text)
    
    except Exception as e:
        print(f"DEBUG: Exception in PDF processing: {e}")
        raise Exception(f"Failed to process PDF: {str(e)}")

def extract_text_from_docx(uploaded_file) -> str:
    """
    Extract text from DOCX file using python-docx
    
    Args:
        uploaded_file: Streamlit uploaded file object
    
    Returns:
        str: Extracted text content
    
    Raises:
        ValueError: If python-docx is not available
    """
    try:
        from docx import Document
    except ImportError:
        raise ValueError("python-docx library is required for DOCX processing. Please install it with: pip install python-docx")
    
    try:
        # Create a Document object
        doc = Document(io.BytesIO(uploaded_file.read()))
        
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        full_text = '\n'.join(text_content)
        
        if not full_text.strip():
            raise Exception("No text could be extracted from the DOCX file.")
        
        return clean_text(full_text)
    
    except ImportError:
        raise ValueError("python-docx library is required for DOCX processing. Please install it with: pip install python-docx")
    except Exception as e:
        if "python-docx" in str(e):
            raise e
        else:
            raise Exception(f"Failed to process DOCX: {str(e)}")

def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text
    
    Args:
        text (str): Raw extracted text
    
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters that might interfere with processing
    text = re.sub(r'[^\w\s\-.,;:()\[\]@#%&*+=<>?/\\|"\'`~]', ' ', text)
    
    # Remove multiple consecutive spaces
    text = re.sub(r' +', ' ', text)
    
    # Strip leading/trailing whitespace
    text = text.strip()
    
    return text

def get_file_info(uploaded_file) -> dict:
    """
    Get information about the uploaded file
    
    Args:
        uploaded_file: Streamlit uploaded file object
    
    Returns:
        dict: File information
    """
    if uploaded_file is None:
        return {}
    
    return {
        'name': uploaded_file.name,
        'size': uploaded_file.size,
        'type': uploaded_file.type,
        'extension': uploaded_file.name.lower().split('.')[-1] if '.' in uploaded_file.name else 'unknown'
    }